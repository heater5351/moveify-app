"""
build_gp_reassessment_report.py — run ONCE to (re)produce
backend/assets/GP_Reassessment_Template.docx

The GP-facing reassessment letter, laid out with the EXACT same chrome as the
initial-consult GP report (gp_report_kit.py): two-tone header, navy section
bands, teal-bordered signature box, navy footer band, Calibri. Only the body
content differs — a baseline-vs-latest comparison instead of the initial
objective assessment, plus a clinical-interpretation + recommendations close.

Unlike build_gp_report.py (which targets the python docxtpl filler), this
template is filled at runtime by backend/services/scribe-gp-reassessment-docx.js
using **docxtemplater**, so the loop uses docxtemplater section syntax
({{#comparison_rows}} … {{/comparison_rows}}), matching the live
GP_Report_Template.docx asset rather than docxtpl's {%tr%} tags.

Tokens consumed (all supplied by scribe-gp-reassessment-docx.js):
  report_date, gp_name (surname), practice_name, practice_address,
  patient_full_name, patient_dob, baseline_date, latest_date,
  cover_letter, executive_summary, clinical_interpretation, recommendations,
  comparison_rows[] {test, baseline, latest, change, interpretation},
  clinician_full_name, clinician_qualifications, clinician_profession,
  clinician_phone, clinician_email, clinician_abn.

Run: python scripts/build_gp_reassessment_report.py
"""

from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

from gp_report_kit import (
    INNER_L,
    NAVY, TEAL, DKGREY, GREY, WHITE,
    set_margins, cell_bg, light_inner_borders, pad, vcenter, sp, rn,
    blank, indented, section_band, page_header, signature_box, footer_band,
)

OUT_PATH = r"C:\Users\dilig\Documents\moveify-app\backend\assets\GP_Reassessment_Template.docx"

SUBTITLE = 'REASSESSMENT REPORT'

doc = Document()
set_margins(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 1 — COVER LETTER
# ══════════════════════════════════════════════════════════════════════════════
page_header(doc, subtitle=SUBTITLE)
blank(doc, 8)

# GP address block (gp_name holds the surname for the reassessment flow).
for txt, bold in [
    ('Dr {{ gp_name }}',       True),
    ('{{ practice_name }}',    False),
    ('{{ practice_address }}', False),
    ('{{ report_date }}',      False),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = INNER_L
    rn(p, txt, size=12, bold=bold, color=NAVY if bold else DKGREY)
    sp(p, 0, 2)

blank(doc, 10)

p = doc.add_paragraph()
p.paragraph_format.left_indent = INNER_L
rn(p, 'Dear Dr {{ gp_name }},', size=12, bold=True, color=NAVY)
sp(p, 0, 10)

# Editable transmittal body. Newlines become line breaks via docxtemplater's
# linebreaks:true, so the multi-paragraph cover letter renders intact.
indented(doc, '{{ cover_letter }}')

blank(doc, 6)

p = doc.add_paragraph()
p.paragraph_format.left_indent = INNER_L
rn(p, 'Yours sincerely,', size=12, color=DKGREY)
# Tighter than the initial-consult report's 38pt: the editable cover letter runs
# a paragraph longer, so this keeps page 1 a single page (no blank spill).
sp(p, 0, 20)

for txt, bold in [
    ('{{ clinician_full_name }}',                                    True),
    ('{{ clinician_qualifications }}  |  {{ clinician_profession }}', False),
    ('{{ clinician_phone }}  |  {{ clinician_email }}',              False),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = INNER_L
    rn(p, txt, size=12, bold=bold, color=NAVY if bold else DKGREY)
    sp(p, 0, 2)

blank(doc, 6)

# Signature box — teal left border
signature_box(doc)

blank(doc, 8)
footer_band(doc)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 2 — REPORT
# ══════════════════════════════════════════════════════════════════════════════
page_header(doc, subtitle=SUBTITLE)
blank(doc, 10)


# ─── PATIENT DETAILS ──────────────────────────────────────────────────────────
section_band(doc, 'Patient Details')
blank(doc, 4)

det = doc.add_table(rows=3, cols=4)
light_inner_borders(det)
det.alignment = WD_TABLE_ALIGNMENT.CENTER
DW = [Cm(3.4), Cm(6.8), Cm(3.4), Cm(6.8)]

det_data = [
    ('Patient Name',  '{{ patient_full_name }}', 'Date of Birth',      '{{ patient_dob }}'),
    ('Referring GP',  'Dr {{ gp_name }}',         'Practice',           '{{ practice_name }}'),
    ('Baseline Date', '{{ baseline_date }}',      'Reassessment Date',  '{{ latest_date }}'),
]
for ri, (l1, v1, l2, v2) in enumerate(det_data):
    row = det.rows[ri]
    bg  = 'F0FAFA' if ri % 2 == 0 else 'FAFEFE'
    for ci, (txt, is_lbl) in enumerate([(l1, True), (v1, False), (l2, True), (v2, False)]):
        c = row.cells[ci]
        c.width = DW[ci]
        cell_bg(c, '1C2E3D' if is_lbl else bg)
        pad(c, top=150, bot=150, left=220, right=160)
        vcenter(c)
        p = c.paragraphs[0]
        p.clear()
        rn(p, txt, size=11, bold=is_lbl, color=WHITE if is_lbl else DKGREY)
        sp(p, 0, 0)

blank(doc, 12)


# ─── EXECUTIVE SUMMARY ────────────────────────────────────────────────────────
section_band(doc, 'Executive Summary')
blank(doc, 5)
indented(doc, '{{ executive_summary }}')
blank(doc, 10)


# ─── OBJECTIVE FINDINGS — BASELINE VS LATEST ──────────────────────────────────
section_band(doc, 'Objective Findings — Baseline vs Latest')
blank(doc, 4)

CW = [Cm(4.6), Cm(2.7), Cm(2.7), Cm(3.0), Cm(7.4)]
cmp = doc.add_table(rows=2, cols=5)
light_inner_borders(cmp, color='C5E5E5')
cmp.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for ci, (h, w) in enumerate(zip(['Measure', 'Baseline', 'Latest', 'Change', 'Clinical Interpretation'], CW)):
    c = cmp.rows[0].cells[ci]
    c.width = w
    cell_bg(c, '1C2E3D')
    pad(c, top=160, bot=160, left=220, right=220)
    p = c.paragraphs[0]
    p.clear()
    rn(p, h, size=11, bold=True, color=WHITE)
    sp(p, 0, 0)

# Template row — docxtemplater repeats this row once per item in comparison_rows.
# {{#comparison_rows}} opens the section in cell 0; {{/comparison_rows}} closes it
# in cell 4. paragraphLoop:true in the filler makes the whole <w:tr> repeat.
loop_row = cmp.rows[1]
loop_cells = [
    ('{{#comparison_rows}}{{ test }}',         True,  False, NAVY),
    ('{{ baseline }}',                          False, False, DKGREY),
    ('{{ latest }}',                            True,  False, NAVY),
    ('{{ change }}',                            True,  False, NAVY),
    ('{{ interpretation }}{{/comparison_rows}}', False, True,  GREY),
]
for ci, (txt, bold, italic, color) in enumerate(loop_cells):
    c = loop_row.cells[ci]
    c.width = CW[ci]
    cell_bg(c, 'FFFFFF')
    pad(c, top=140, bot=140, left=220, right=220)
    vcenter(c)
    p = c.paragraphs[0]
    p.clear()
    rn(p, txt, size=11, bold=bold, italic=italic, color=color)
    sp(p, 0, 0)

blank(doc, 12)


# ─── CLINICAL INTERPRETATION ──────────────────────────────────────────────────
section_band(doc, 'Clinical Interpretation')
blank(doc, 5)
indented(doc, '{{ clinical_interpretation }}')
blank(doc, 12)


# ─── RECOMMENDATIONS ──────────────────────────────────────────────────────────
section_band(doc, 'Recommendations')
blank(doc, 5)
indented(doc, '{{ recommendations }}')
blank(doc, 20)


# ─── FOOTER ───────────────────────────────────────────────────────────────────
footer_band(doc)


# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Template saved: {OUT_PATH}")
