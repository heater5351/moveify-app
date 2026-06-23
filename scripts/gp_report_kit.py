"""
gp_report_kit.py — shared Moveify GP-report chrome (python-docx).

Single source of truth for the GP-facing letter/report look:
  · two-tone page header (logo + navy subtitle box + teal accent bar)
  · full-width navy section bands
  · light teal inner-border tables
  · teal-bordered signature box
  · navy footer band with clinician line

Imported by:
  - build_gp_report.py             → gp_report_template.docx  (initial consult)
  - build_gp_reassessment_report.py → backend/assets/GP_Reassessment_Template.docx

Keep ALL styling here so the two reports stay visually identical — restyle once.

Note on placeholder syntax: every {{ variable }} / loop tag is written into a
SINGLE run by rn(), so the docxtemplater filler in the backend can find and
replace them reliably without run-splitting.
"""

from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL  # noqa: F401 (re-exported for callers)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Asset paths ─────────────────────────────────────────────────────────────────
LOGO_PATH = r"C:\Users\dilig\Documents\moveify-app\frontend\public\assets\moveify-logo.png"

# ── Page geometry ────────────────────────────────────────────────────────────────
INNER_L = Cm(2.0)
INNER_R = Cm(2.0)

# ── Colours ──────────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x13, 0x22, 0x32)
TEAL   = RGBColor(0x46, 0xC1, 0xC0)
DKGREY = RGBColor(0x3A, 0x44, 0x52)
GREY   = RGBColor(0x9C, 0xA3, 0xAF)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


# ── Page setup ───────────────────────────────────────────────────────────────────
def set_margins(doc):
    sec = doc.sections[0]
    sec.top_margin    = Cm(0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(0)
    sec.right_margin  = Cm(0)


# ── Low-level helpers ────────────────────────────────────────────────────────────
def cell_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)


def no_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    b = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        b.append(el)
    tblPr.append(b)


def light_inner_borders(table, color='CBE9E9'):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    b = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        b.append(el)
    for side in ('insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        b.append(el)
    tblPr.append(b)


def pad(cell, top=100, bot=100, left=160, right=160):
    tcPr  = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, v in [('top', top), ('bottom', bot), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(v))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def vcenter(cell):
    tcPr   = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def sp(p, before=0, after=5):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)


def rn(p, text, size=12, bold=False, italic=False, color=None):
    """Add a single run — all text in one run so placeholders are never split."""
    r = p.add_run(text)
    r.font.name   = 'Calibri'
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def blank(doc, pts=8):
    p = doc.add_paragraph()
    sp(p, 0, pts)


def indented(doc, text, size=12, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = INNER_L
    p.paragraph_format.right_indent = INNER_R
    p.paragraph_format.line_spacing = Pt(17)
    rn(p, text, size=size, italic=italic, color=color or DKGREY)
    sp(p, 0, 6)
    return p


# ── Chrome blocks ────────────────────────────────────────────────────────────────
def section_band(doc, title):
    t = doc.add_table(rows=1, cols=1)
    no_borders(t)
    c = t.cell(0, 0)
    cell_bg(c, '132232')
    pad(c, top=200, bot=200, left=560, right=560)
    p = c.paragraphs[0]
    p.clear()
    rn(p, title.upper(), size=11, bold=True, color=WHITE)
    sp(p, 0, 0)


def page_header(doc, subtitle='INITIAL CONSULTATION REPORT'):
    hdr = doc.add_table(rows=1, cols=2)
    no_borders(hdr)
    lc = hdr.cell(0, 0)
    rc = hdr.cell(0, 1)
    lc.width = Cm(13.5)
    rc.width = Cm(7.5)
    cell_bg(lc, 'FFFFFF')
    cell_bg(rc, '132232')
    pad(lc, top=200, bot=120, left=480, right=200)
    pad(rc, top=200, bot=120, left=300, right=400)
    vcenter(lc)
    vcenter(rc)

    lp = lc.paragraphs[0]
    lp.clear()
    run = lp.add_run()
    run.add_picture(LOGO_PATH, width=Cm(7.0))
    sp(lp, 0, 0)

    rp = rc.paragraphs[0]
    rp.clear()
    rp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rn(rp, subtitle, size=11, bold=True, color=WHITE)
    sp(rp, 0, 6)

    rp2 = rc.add_paragraph()
    rp2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rn(rp2, 'Exercise Physiology  ·  Allied Health', size=9, color=TEAL)
    sp(rp2, 0, 0)

    acc = doc.add_table(rows=1, cols=1)
    no_borders(acc)
    ac = acc.cell(0, 0)
    cell_bg(ac, '46C1C0')
    pad(ac, top=60, bot=60, left=0, right=0)
    ac.paragraphs[0].clear()
    sp(ac.paragraphs[0], 0, 0)


def signature_box(doc, label='[ Clinician Signature ]'):
    """Teal-left-bordered signature panel, indented to align under the letter body."""
    sig_t = doc.add_table(rows=1, cols=1)
    no_borders(sig_t)
    tblPr = sig_t._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        sig_t._tbl.insert(0, tblPr)
    ind_el = OxmlElement('w:tblInd')
    ind_el.set(qn('w:w'), '1134')
    ind_el.set(qn('w:type'), 'dxa')
    tblPr.append(ind_el)

    sc = sig_t.cell(0, 0)
    sc.width = Cm(7)
    cell_bg(sc, 'EBF8F8')
    pad(sc, top=320, bot=320, left=300, right=300)

    tcPr = sc._tc.get_or_add_tcPr()
    tcb  = OxmlElement('w:tcBorders')
    lbrd = OxmlElement('w:left')
    lbrd.set(qn('w:val'),   'single')
    lbrd.set(qn('w:sz'),    '20')
    lbrd.set(qn('w:space'), '0')
    lbrd.set(qn('w:color'), '46C1C0')
    tcb.append(lbrd)
    tcPr.append(tcb)

    p = sc.paragraphs[0]
    p.clear()
    rn(p, label, size=11, italic=True, color=GREY)
    sp(p, 0, 0)


def footer_band(doc):
    ft = doc.add_table(rows=1, cols=1)
    no_borders(ft)
    fc = ft.cell(0, 0)
    cell_bg(fc, '132232')
    pad(fc, top=160, bot=160, left=560, right=560)
    p = fc.paragraphs[0]
    p.clear()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(p, 'Moveify Health Solutions  ·  Exercise Physiology  ·  Allied Health', size=9, color=TEAL)
    sp(p, 0, 3)
    p2 = fc.add_paragraph()
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(p2, '{{ clinician_full_name }}  |  {{ clinician_qualifications }}  |  {{ clinician_phone }}  |  {{ clinician_email }}  |  ABN: {{ clinician_abn }}',
       size=9, color=GREY)
    sp(p2, 0, 0)
