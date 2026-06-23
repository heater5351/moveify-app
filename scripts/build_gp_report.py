"""
build_gp_report.py — run ONCE to produce gp_report_template.docx

The output is a docxtpl-compatible Word template.  Every {{ variable }} and
{%tr for %} tag is written as plain text by python-docx into a single run, so
docxtpl can find and replace them reliably at fill-time.

Two-step workflow:
  1. python scripts/build_gp_report.py   →  produces gp_report_template.docx
  2. python scripts/fill_gp_report.py    →  fills template with patient data + AI content

Chrome (header / section bands / signature box / footer) lives in gp_report_kit.py
and is shared with build_gp_reassessment_report.py so the two reports stay identical.
"""

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from gp_report_kit import (
    INNER_L, INNER_R,
    NAVY, TEAL, DKGREY, GREY, WHITE,
    set_margins, cell_bg, no_borders, light_inner_borders, pad, vcenter, sp, rn,
    blank, indented, section_band, page_header, signature_box, footer_band,
)

OUT_PATH = r"C:\Users\dilig\Documents\moveify-app\gp_report_template.docx"

doc = Document()
set_margins(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 1 — COVER LETTER
# ══════════════════════════════════════════════════════════════════════════════
page_header(doc)
blank(doc, 14)

# GP address block
for txt, bold in [
    ('{{ gp_name }}',        True),
    ('{{ practice_name }}',  False),
    ('{{ practice_address }}', False),
    ('{{ report_date }}',    False),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = INNER_L
    rn(p, txt, size=12, bold=bold, color=NAVY if bold else DKGREY)
    sp(p, 0, 2)

blank(doc, 10)

p = doc.add_paragraph()
p.paragraph_format.left_indent = INNER_L
rn(p, 'Dear Dr {{ gp_surname }},', size=12, bold=True, color=NAVY)
sp(p, 0, 10)

indented(doc,
    'Thank you sincerely for referring {{ patient_full_name }} to Moveify Health Solutions for '
    'Exercise Physiology services under the Chronic Disease Management (CDM) Plan. Please find '
    'below the report and recommendations following {{ patient_pronoun }} Initial Consultation '
    'on {{ appointment_date }}.')

indented(doc,
    'Should you have any questions or queries, please do not hesitate to contact me on '
    '{{ clinician_phone }}.')

blank(doc, 10)

p = doc.add_paragraph()
p.paragraph_format.left_indent = INNER_L
rn(p, 'Yours sincerely,', size=12, color=DKGREY)
sp(p, 0, 38)

for txt, bold in [
    ('{{ clinician_full_name }}',                             True),
    ('{{ clinician_qualifications }}  |  {{ clinician_profession }}', False),
    ('{{ clinician_phone }}  |  {{ clinician_email }}',       False),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = INNER_L
    rn(p, txt, size=12, bold=bold, color=NAVY if bold else DKGREY)
    sp(p, 0, 2)

blank(doc, 12)

# Signature box — teal left border
signature_box(doc)

blank(doc, 16)
footer_band(doc)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 2 — REPORT
# ══════════════════════════════════════════════════════════════════════════════
page_header(doc)
blank(doc, 10)


# ─── PATIENT DETAILS ──────────────────────────────────────────────────────────
section_band(doc, 'Patient Details')
blank(doc, 4)

det = doc.add_table(rows=3, cols=4)
light_inner_borders(det)
det.alignment = WD_TABLE_ALIGNMENT.CENTER
DW = [Cm(3.0), Cm(7.2), Cm(3.0), Cm(7.2)]

det_data = [
    ('Patient Name',  '{{ patient_full_name }}', 'Date of Birth',  '{{ patient_dob }}'),
    ('Referring GP',  '{{ referring_gp }}',       'Practice',       '{{ practice_name }}'),
    ('Referral Date', '{{ referral_date }}',       'CDM Sessions',   '{{ cdm_sessions }}'),
]
for ri, (l1, v1, l2, v2) in enumerate(det_data):
    row = det.rows[ri]
    bg  = 'F0FAFA' if ri % 2 == 0 else 'FAFEFE'
    for ci, (txt, is_lbl) in enumerate([(l1, True), (v1, False), (l2, True), (v2, False)]):
        c = row.cells[ci]
        c.width = DW[ci]
        cell_bg(c, '1C2E3D' if is_lbl else bg)
        pad(c, top=150, bot=150, left=220, right=160)
        vcenter(c)
        p = c.paragraphs[0]
        p.clear()
        rn(p, txt, size=11, bold=is_lbl, color=WHITE if is_lbl else DKGREY)
        sp(p, 0, 0)

blank(doc, 12)


# ─── EXECUTIVE SUMMARY ────────────────────────────────────────────────────────
section_band(doc, 'Executive Summary')
blank(doc, 5)
indented(doc, '{{ executive_summary }}')
blank(doc, 10)


# ─── OBJECTIVE ASSESSMENT ─────────────────────────────────────────────────────
section_band(doc, 'Objective Assessment')
blank(doc, 4)

OA_W = [Cm(5.4), Cm(3.0), Cm(12.0)]
oa = doc.add_table(rows=2, cols=3)
light_inner_borders(oa, color='C5E5E5')
oa.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for ci, (h, w) in enumerate(zip(['Test', 'Result', 'Interpretation'], OA_W)):
    c = oa.rows[0].cells[ci]
    c.width = w
    cell_bg(c, '1C2E3D')
    pad(c, top=160, bot=160, left=220, right=220)
    p = c.paragraphs[0]
    p.clear()
    rn(p, h, size=11, bold=True, color=WHITE)
    sp(p, 0, 0)

# Template row — docxtpl loops over oa_rows and repeats this row for each item.
# {%tr for row in oa_rows %} opens the loop in cell 0;
# {%tr endfor %}            closes it in cell 2.
loop_row = oa.rows[1]
loop_cells = [
    ('{%tr for row in oa_rows %}{{ row.test }}', False, NAVY),
    ('{{ row.result }}',                          False, DKGREY),
    ('{{ row.interpretation }}{%tr endfor %}',    True,  GREY),
]
for ci, (txt, italic, color) in enumerate(loop_cells):
    c = loop_row.cells[ci]
    c.width = OA_W[ci]
    cell_bg(c, 'FFFFFF')
    pad(c, top=140, bot=140, left=220, right=220)
    vcenter(c)
    p = c.paragraphs[0]
    p.clear()
    rn(p, txt, size=11, bold=(ci == 0), italic=italic, color=color)
    sp(p, 0, 0)

blank(doc, 12)


# ─── GOALS ────────────────────────────────────────────────────────────────────
section_band(doc, 'Goals')
blank(doc, 5)
indented(doc, '{{ goals }}')
blank(doc, 12)


# ─── MANAGEMENT PLAN ──────────────────────────────────────────────────────────
section_band(doc, 'Management Plan')
blank(doc, 5)
indented(doc, '{{ management_plan }}')
blank(doc, 20)


# ─── FOOTER ───────────────────────────────────────────────────────────────────
footer_band(doc)


# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Template saved: {OUT_PATH}")
print()
print("Next step: run  python scripts/fill_gp_report.py  to generate a filled report.")
